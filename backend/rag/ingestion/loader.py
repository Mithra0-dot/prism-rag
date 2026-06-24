"""
PRISM — Precision Retrieval with Intelligent Semantic Multimodal
----------------------------------------------------------------
backend/rag/ingestion/loader.py

Document loading and parsing pipeline — Phase 2 upgrade.
Now supports vision=True for OCR extraction from images/charts in PDFs.
"""

import time
from pathlib import Path
from datetime import datetime
from typing import Optional

from langchain.schema import Document

from backend.core.config import get_settings
from backend.core.logger import get_logger
from backend.core.exceptions import (
    UnsupportedFileTypeError,
    FileTooLargeError,
    DocumentIngestionError,
)

settings = get_settings()
log = get_logger("loader")


class DocumentLoader:
    """
    Unified document loader for PRISM.

    Phase 2 upgrade: pass vision=True to activate OCR extraction
    for images and charts embedded in PDFs.
    """

    _PARSER_MAP: dict[str, str] = {
        ".pdf":  "_load_pdf",
        ".txt":  "_load_text",
        ".md":   "_load_text",
        ".docx": "_load_docx",
    }

    def __init__(self) -> None:
        self._allowed_extensions = set(settings.allowed_extensions)
        self._max_bytes = settings.max_upload_size_bytes
        log.info(
            f"DocumentLoader initialised | "
            f"allowed={self._allowed_extensions} | "
            f"max_size={settings.max_upload_size_mb}MB"
        )

    async def load(
        self,
        file_path: Path,
        filename: Optional[str] = None,
        vision: bool = False,
    ) -> list[Document]:
        """
        Load a document and return a list of Document objects.

        Args:
            file_path : Path to the file on disk
            filename  : Original filename for metadata
            vision    : If True and file is PDF, run OCR on embedded images

        Returns:
            List of LangChain Document objects
        """
        filename = filename or file_path.name
        extension = file_path.suffix.lower()

        log.info(f"Loading document: {filename} ({extension}) | vision={vision}")
        start_time = time.perf_counter()

        self._validate_extension(extension, filename)
        self._validate_size(file_path, filename)

        try:
            # Use vision pipeline for PDFs when requested
            if vision and extension == ".pdf":
                documents = self._load_pdf_with_vision(file_path, filename)
            else:
                parser_method = getattr(self, self._PARSER_MAP[extension])
                documents = parser_method(file_path, filename)
        except (UnsupportedFileTypeError, FileTooLargeError):
            raise
        except Exception as e:
            raise DocumentIngestionError(
                message=f"Failed to parse '{filename}'",
                detail=str(e),
            ) from e

        documents = self._filter_empty(documents)
        documents = self._attach_index(documents)

        elapsed = time.perf_counter() - start_time
        log.info(
            f"Loaded '{filename}' → {len(documents)} pages/sections "
            f"in {elapsed:.2f}s"
        )

        return documents

    def load_sync(
        self,
        file_path: Path,
        filename: Optional[str] = None,
        vision: bool = False,
    ) -> list[Document]:
        import asyncio
        return asyncio.run(self.load(file_path, filename, vision=vision))

    # ── Validators ────────────────────────────────────────────────────

    def _validate_extension(self, extension: str, filename: str) -> None:
        if extension not in self._allowed_extensions:
            raise UnsupportedFileTypeError(
                message=f"'{extension}' files are not supported by PRISM",
                detail=f"filename='{filename}' | allowed={sorted(self._allowed_extensions)}",
            )

    def _validate_size(self, file_path: Path, filename: str) -> None:
        size_bytes = file_path.stat().st_size
        if size_bytes > self._max_bytes:
            size_mb = size_bytes / (1024 * 1024)
            raise FileTooLargeError(
                message=f"'{filename}' is {size_mb:.1f}MB — limit is {settings.max_upload_size_mb}MB",
                detail=f"size_bytes={size_bytes}",
            )

    # ── Parsers ───────────────────────────────────────────────────────

    def _load_pdf(self, file_path: Path, filename: str) -> list[Document]:
        """Standard text-only PDF loading via pypdf."""
        try:
            from pypdf import PdfReader
        except ImportError as e:
            raise DocumentIngestionError(message="pypdf is not installed", detail=str(e))

        reader = PdfReader(str(file_path))
        documents: list[Document] = []

        for page_num, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            documents.append(Document(
                page_content=text,
                metadata=self._build_metadata(
                    filename=filename,
                    file_type="pdf",
                    page=page_num,
                    total_pages=len(reader.pages),
                ),
            ))

        return documents

    def _load_pdf_with_vision(self, file_path: Path, filename: str) -> list[Document]:
        """
        Phase 2: PDF loading with OCR for embedded images and charts.
        Falls back to standard loading if vision dependencies are missing.
        """
        try:
            from backend.rag.ingestion.vision import VisionExtractor
            extractor = VisionExtractor()
            log.info(f"Running vision extraction on '{filename}'")
            return extractor.extract(pdf_path=file_path, filename=filename)
        except ImportError:
            log.warning(
                "Vision dependencies not available — falling back to text-only. "
                "Run: pip install pymupdf easyocr opencv-python-headless"
            )
            return self._load_pdf(file_path, filename)
        except Exception as e:
            log.warning(f"Vision extraction failed, falling back to text: {e}")
            return self._load_pdf(file_path, filename)

    def _load_text(self, file_path: Path, filename: str) -> list[Document]:
        """Parse plain text and Markdown files."""
        extension = file_path.suffix.lower()
        file_type = "markdown" if extension == ".md" else "text"

        try:
            text = file_path.read_text(encoding="utf-8", errors="replace").strip()
        except OSError as e:
            raise DocumentIngestionError(
                message=f"Could not read '{filename}'", detail=str(e)
            )

        if not text:
            return []

        return [Document(
            page_content=text,
            metadata=self._build_metadata(
                filename=filename, file_type=file_type, page=1, total_pages=1
            ),
        )]

    def _load_docx(self, file_path: Path, filename: str) -> list[Document]:
        """Parse DOCX files using python-docx."""
        try:
            import docx
        except ImportError as e:
            raise DocumentIngestionError(
                message="python-docx not installed. Run: pip install python-docx",
                detail=str(e),
            )

        doc = docx.Document(str(file_path))
        documents: list[Document] = []
        current_section: list[str] = []
        section_num = 1

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            is_heading = para.style.name.startswith("Heading")
            if is_heading and current_section:
                documents.append(Document(
                    page_content="\n".join(current_section),
                    metadata=self._build_metadata(
                        filename=filename, file_type="docx",
                        page=section_num, total_pages=None,
                    ),
                ))
                current_section = []
                section_num += 1
            current_section.append(text)

        if current_section:
            documents.append(Document(
                page_content="\n".join(current_section),
                metadata=self._build_metadata(
                    filename=filename, file_type="docx",
                    page=section_num, total_pages=None,
                ),
            ))

        return documents

    # ── Helpers ───────────────────────────────────────────────────────

    def _build_metadata(
        self,
        filename: str,
        file_type: str,
        page: int,
        total_pages: Optional[int],
    ) -> dict:
        meta = {
            "source":      filename,
            "file_type":   file_type,
            "page":        page,
            "ingested_at": datetime.utcnow().isoformat(),
            "pipeline":    "text",
        }
        if total_pages is not None:
            meta["total_pages"] = total_pages
        return meta

    def _filter_empty(self, documents: list[Document]) -> list[Document]:
        filtered = [d for d in documents if d.page_content.strip()]
        removed = len(documents) - len(filtered)
        if removed:
            log.debug(f"Filtered out {removed} empty document sections")
        return filtered

    def _attach_index(self, documents: list[Document]) -> list[Document]:
        for i, doc in enumerate(documents):
            doc.metadata["chunk_index"] = i
        return documents